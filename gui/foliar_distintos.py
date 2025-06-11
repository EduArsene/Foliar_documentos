from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils.docx_tools import agregar_elemento, configurar_idioma_parrafo

def impar(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+ 1')  # + 1
    agregar_elemento(correr, 'w:fldChar', 'separate')      # separador
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+ 1')  # + 1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque      
#para poder ponerlo en el Encabezado
def foliarImpar(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
# para poder ponerlo en el Pie
def foliarImparPie(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#**********************************************  
#para foliar PAR en PIE y ENCABEZADO de pagina    Insercion de codigo de capo 
def impar_Desc(parrafo,numPaginas):
    correr = parrafo.add_run()
    paginas = numPaginas
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES') # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'separate')      # separador
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    if paginas% 2 !=0 :
        agregar_elemento(correr, 'w:instrText', 'begin', '/2+0.5') # /2
    else: agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES') # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    if paginas% 2 !=0 :
        agregar_elemento(correr, 'w:instrText', 'begin', '/2+0.5') # /2
    else: agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
#Foliacion impar de manera descendente en el Encabezado
def foliarImparDesc(doc,numPaginas):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar_Desc(parrafo,numPaginas)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#Foliacion impar de manera descendente en el Pie de pagina
def foliarImparPieDesc(doc,numPaginas):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar_Desc(parrafo,numPaginas)
    configurar_idioma_parrafo(parrafo, 'es-PE')
    
#**********************************************   
#Funcoin para foliacion normal   Insercion de codigo de campo 
def Normal(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }    
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque

#Insertar de manera Normal en el Encabezado
def foliarNormal(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
    
#Insertar de manera Normal en el Pie de pagina
def foliarNormalPie(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#**********************************************  
#Foliacion normal de manera descente     Insercion de codigo de capo en word
def Normal_Desc(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque { = { NUMPAGES } - { PAGE } + 1 }
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES')  # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque {={={numpages}-{page}+1}\*\*}
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES')  # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }   
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque
def foliarNormalDesc(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal_Desc(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
      
def foliarNormalPieDesc(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal_Desc(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
