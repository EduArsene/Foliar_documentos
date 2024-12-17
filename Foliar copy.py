#Importar librerias
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import ttk
import os
from tkinter import messagebox, filedialog
from PIL import Image, ImageTk
import importlib.resources as pkg_resources
import src  # Importar el paquete src

#**********************************************
# Función para ocultar la consola
def hide_console():
    import ctypes
    hwnd = ctypes.windll.kernel32.GetConsoleWindow()
    if hwnd != 0:
        ctypes.windll.user32.ShowWindow(hwnd, 0)  # Ocultar la consola

#**********************************************          
#Para agregar codigo de campo
def agregar_elemento(agregar, tag, tipo, texto=None):
    elemento = OxmlElement(tag)
    elemento.set(qn('w:fldCharType'), tipo)
    if texto:
        elemento.set(qn('xml:space'), 'preserve')
        elemento.text = texto
    agregar._r.append(elemento)
#************************************   
#Para poder activar opcion de headres diferentes
def activar_pares_impares_diferentes(seccion):
    seccPr = seccion._sectPr
    if seccPr is None:
        seccPr = OxmlElement('w:sectPr')
        seccion._element.append(seccPr)
    
    # Eliminar la opción de "Primera página diferente" si existe
    for el in seccPr.findall(qn('w:titlePg')):
        seccPr.remove(el)
    
    # Activar la opción de "Páginas pares e impares diferentes"
    evenAndOddHeaders = OxmlElement('w:evenAndOddHeaders')
    seccPr.append(evenAndOddHeaders)
    
#**********************************************  
#*************configuracion de idioma   
def configurar_idioma_parrafo(parrafo, lang_code):
    for run in parrafo.runs:
        rPr = run._element.get_or_add_rPr()
        lang = OxmlElement('w:lang')
        lang.set(qn('w:val'), lang_code)
        rPr.append(lang)

#************************************   
#funcion de foliacion IMPAR en ENCABEZADO y PIE de PAGINA   Insercion de codigo de capo
def impar(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+ 1')  # + 1
    agregar_elemento(correr, 'w:fldChar', 'separate')      # separador
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+ 1')  # + 1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque      
#para poder ponerlo en el Encabezado
def foliarImpar(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
# para poder ponerlo en el Pie
def foliarImparPie(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#**********************************************  
#para foliar PAR en PIE y ENCABEZADO de pagina    Insercion de codigo de capo 
def impar_Desc(parrafo,numPaginas):
    correr = parrafo.add_run()
    paginas = numPaginas
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES') # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'separate')      # separador
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    if paginas% 2 !=0 :
        agregar_elemento(correr, 'w:instrText', 'begin', '/2+0.5') # /2
    else: agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES') # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE') # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    if paginas% 2 !=0 :
        agregar_elemento(correr, 'w:instrText', 'begin', '/2+0.5') # /2
    else: agregar_elemento(correr, 'w:instrText', 'begin', '/2') # /2
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
#Foliacion impar de manera descendente en el Encabezado
def foliarImparDesc(doc,numPaginas):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar_Desc(parrafo,numPaginas)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#Foliacion impar de manera descendente en el Pie de pagina
def foliarImparPieDesc(doc,numPaginas):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    impar_Desc(parrafo,numPaginas)
    configurar_idioma_parrafo(parrafo, 'es-PE')
    
#**********************************************   
#Funcoin para foliacion normal   Insercion de codigo de campo 
def Normal(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # }    
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque

#Insertar de manera Normal en el Encabezado
def foliarNormal(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
    
#Insertar de manera Normal en el Pie de pagina
def foliarNormalPie(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#**********************************************  
#Foliacion normal de manera descente     Insercion de codigo de capo en word
def Normal_Desc(parrafo):
    correr = parrafo.add_run()
    # Inicio primer bloque { = { NUMPAGES } - { PAGE } + 1 }
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES')  # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #fin primer bloque
    agregar_elemento(correr, 'w:instrText', 'begin', ': ')  # :
    #Inicio segundo bloque {={={numpages}-{page}+1}\*\*}
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', '=')  # =
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'NUMPAGES')  # NUMPAGES
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '-')  # -
    agregar_elemento(correr, 'w:fldChar', 'begin')         # {
    agregar_elemento(correr, 'w:instrText', 'begin', 'PAGE')  # PAGE
    agregar_elemento(correr, 'w:fldChar', 'end')           # } 
    agregar_elemento(correr, 'w:instrText', 'begin', '+1')  # +1
    agregar_elemento(correr, 'w:fldChar', 'end')           # }   
    agregar_elemento(correr, 'w:instrText', 'begin', '\*CardText \*Upper') # \*CardText \*Upper
    agregar_elemento(correr, 'w:fldChar', 'end')           # }
    #Fin segundo bloque
def foliarNormalDesc(doc):
    seccion = doc.sections[0]
    header = seccion.header
    parrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal_Desc(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')
      
def foliarNormalPieDesc(doc):
    seccion = doc.sections[0]
    footer = seccion.footer
    parrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    Normal_Desc(parrafo)
    configurar_idioma_parrafo(parrafo, 'es-PE')

#**********************************************  
#Abrri directorio para poder visualizar el archivo creado
def abrir_directorio(archivo):
        directorio = os.path.dirname(archivo)
        if os.path.exists(directorio):
            os.startfile(directorio)
        else:
            messagebox.showerror("Error", "No se pudo abrir el directorio.")
#**********************************************            
#Abrir imagenes sin la nacesidad de poner una ruta especifica
def abrir_imagen(nombre_archivo):
    with pkg_resources.path(src, nombre_archivo) as ruta_imagen:

        return ruta_imagen
#**********************************************     
#funcion principal para crear
def crear_documento_foliado():
    import PyPDF2
    from docx import Document
    #minimizar la ventana principal
    root.iconify()
    numero_paginas = tk.IntVar()
    #funcion para poder selccionar el pdf principal
    def seleccionar_pdf():
        try:
            archivo_pdf = filedialog.askopenfilename(filetypes=[("Archivos PDF", "*.pdf")])
            if archivo_pdf:
                ruta_pdf.set(archivo_pdf)
                try:
                    reader = PyPDF2.PdfReader(archivo_pdf)
                    
                    numero_paginas.set(len(reader.pages))
                    estado.set("Archivo cargado")
                    boton_crear['state'] = tk.NORMAL                
                except Exception as e:
                    messagebox.showerror("Error", f"No se pudo leer el archivo PDF: {e}")
        except ValueError as e:
            messagebox.showerror("Error", str(e))
    #Funcion para poder obtener la orientacion del archivo pdf
    def obtener_orientacion(ruta_pdf):
        with open(ruta_pdf, 'rb') as file:
            pdf_leer = PyPDF2.PdfReader(file)
            pagina = pdf_leer.pages[0]
            
            page_ancho = pagina.mediabox[2]
            page_altura = pagina.mediabox[3]
            
            if page_ancho > page_altura:
                return 'Horizontal'
            else:
                return 'Vertical' 
    #Funcion para superponer los archivos           
    def superponer_pdfs_horizontal():
    # Abrir los archivos PDF pdf1_path, pdf2_path
        archivo_abajo_datos = ruta_pdf.get()
        if obtener_orientacion(ruta_pdf.get()) == "Horizontal" and opcion_ori.get() == "Horizontal":
            archivo_arriba = "numeracion_impar_vertical.pdf"
            if opcion_var.get() == "Normal":
                archivo_arriba = "numeracion_normal_vertical.pdf"    
        else:
            archivo_arriba = aux_ruta_guardar.get() 
            
        with open(archivo_abajo_datos, 'rb') as file1, open(archivo_arriba, 'rb') as file2:
            reader1 = PyPDF2.PdfReader(file1)
            reader2 = PyPDF2.PdfReader(file2)
            writer = PyPDF2.PdfWriter()

            # Iterar sobre las páginas del primer PDF
            for page_num in range(len(reader1.pages)):
                page1 = reader1.pages[page_num]

                # Si hay una página correspondiente en el segundo PDF, superponerla
                if page_num < len(reader2.pages):
                    page2 = reader2.pages[page_num]
                    page1.merge_page(page2)  # Superponer la página

                # Agregar la página resultante al escritor
                writer.add_page(page1)
                
                resultado = ruta_pdf.get()
                resultado = resultado.upper()
                resultado = resultado.replace(".PDF", "-FOLIADO.pdf")
            # Guardar el archivo PDF resultante
            with open(resultado, 'wb') as output_file:
                writer.write(output_file)
                
    def crear_doc():
        from docx.enum.section import WD_ORIENT
        from docx2pdf import convert
        from docx.shared import Cm
        from docx.shared import Inches
        from docx.shared import Mm
        # llamar a rutas, la principal y donde se guardará lo creado
        archivo_pdf = ruta_pdf.get()
        archivo_guardar = ruta_guardar.get()
        if not archivo_pdf:
            messagebox.showerror("Error", "Seleccione un archivo PDF válido.")
            return
        if not archivo_guardar:
            # Si no se selecciona una ruta, se guarda en el mismo directorio del archivo PDF original
            archivo_guardar = os.path.join(os.path.dirname(archivo_pdf), "resultado_foliado.pdf")
            aux_ruta_guardar.set(archivo_guardar)  # Actualizar la variable auxiliar con la ruta por defecto     
        try:#obtener el numero de las pagians
            numero_de_paginas = int(numero_paginas.get())
            if numero_de_paginas <= 0:
                raise ValueError("El número de páginas debe ser mayor a 0")
        except ValueError as e:
            messagebox.showerror("Error", str(e))
            return
        #variable para manipular el doc
        doc = Document()
        #agregar numero de la cantidad ed las paginas necesarias
        for i in range(1, numero_de_paginas):
            doc.add_page_break()      
        #guardar la orientacion en una variable   
        orientacion = obtener_orientacion(ruta_pdf.get())
        for section in doc.sections:
            section.page_height = Mm(297)  # Alto A4 en mm
            section.page_width = Mm(210)   # Ancho A4 en mm
            section.header_distance = Cm(0.5) 
            section.footer_distance = Cm(0.85)
            if orientacion == "Horizontal":
                section.right_margin = Inches(0.1)
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width                
            else:
                section.orientation = WD_ORIENT.PORTRAIT      
        try:#condicionales para poder foliar en distintos asos            
            opcion_PE = pie_encabezado.get()        
            opcionIP = opcion_var.get()            
            opcion_AD = opcion_asc_desc.get()           
            if opcion_PE == "Encabezado":#caso Encabezado de pagina
               if opcionIP == "Impar" :
                   if opcion_AD == "Ascendente":foliarImpar(doc)
                   else: foliarImparDesc(doc,numero_de_paginas)
               elif opcionIP == "Normal":
                  if opcion_AD == "Ascendente":foliarNormal(doc)  
                  else: foliarNormalDesc(doc)                                 
            else:                        #caso PIE de pagina
                if opcionIP == "Impar":
                    if opcion_AD == "Ascendente":foliarImparPie(doc)
                    else: foliarImparPieDesc(doc,numero_de_paginas)
                elif opcionIP == "Normal": 
                    if opcion_AD == "Ascendente":foliarNormalPie(doc)  
                    else: foliarNormalPieDesc(doc)                                                               
            if opcionIP != "Normal":
                for section in doc.sections:
                    activar_pares_impares_diferentes(section)
            #creacion de un archivo temporal para superponer            
            archivo_docx_temp = "documento_temporal.docx"
            doc.save(archivo_docx_temp)
            #convertir de Docx a PDF
            convert(archivo_docx_temp, archivo_guardar,True)
            #borrar el temporal 
            os.remove(archivo_docx_temp)
            #llamada a la funcion superponer para poder fusionar los PDF
            superponer_pdfs_horizontal()
            ventana_config.update_idletasks()
            #enviando la ruta del archivo para poder abrir el directorio
            abrir_directorio(archivo_guardar)
            os.remove(archivo_guardar)
        except ValueError as e:
            messagebox.showerror("Error", str(e))
#**********************************************  
#mostrar imagen ejemplo
    def mostrar_ayuda():
        ventana_imagen = tk.Toplevel()
        ventana_imagen.title("MUESTRA DE FOLIACION")
        ventana_imagen.attributes("-topmost", True)
        alto = 1100
        ancho = 700
        ventana_imagen.maxsize(alto,ancho)  
        ventana_imagen.minsize(alto,ancho)
        canvas = tk.Canvas(ventana_imagen, width=alto, height=ancho)#("2424x1593")
        canvas.pack()
        imagen_original = Image.open(abrir_imagen("Muestra.png"))
        
        imagen_redimensionada = imagen_original.resize((alto, ancho))  # Ajusta el tamaño según lo necesario
        imagen = ImageTk.PhotoImage(imagen_redimensionada)
        
        canvas.create_image(0, 0, anchor=tk.NW, image=imagen)
        canvas.image = imagen#largoxalto
        ventana_imagen.geometry(f"{alto}x{ancho}")
    #**********************************************  
    #minimizar la ventana principal
    def minimizar_ventana_principal():
        #minimirzar principal
        root.deiconify()
        ventana_config.destroy()
    #crear ventana 
    ventana_config = tk.Toplevel(root)
    ventana_config.title("Crear Documento Foliado")
    ventana_width = 850  # Ancho de la nueva ventana
    ventana_height = 640  # Altura de la nueva ventana
    ventana_config.geometry(f"{ventana_width}x{ventana_height}")#aumentar +{ventana_w}+{ventana_h} par aponer ne cierto sitio
    ventana_config.minsize(ventana_width,ventana_height)
    ventana_config.maxsize(ventana_width,ventana_height)
    ventana_config.configure(bg=color_secundario)
    ventana_config.protocol("WM_DELETE_WINDOW", minimizar_ventana_principal)
#******************SEUGNDA INTERFAZ******************************
    #Cargar la imagen de fondo
    imagen_fondo2 = Image.open(abrir_imagen("fondo2.jpg"))
    imagen_fondo2 = imagen_fondo2.resize((ventana_width, ventana_height))  # Ajustar tamaño si es necesario
    fondo2 = ImageTk.PhotoImage(imagen_fondo2)
    #Crear un label para mostrar la imagen de fondo
    label_fondo2 = tk.Label(ventana_config, image=fondo2)
    label_fondo2.place(x=0, y=0, relwidth=1, relheight=1)
    
    #franja 
    tit_crear = tk.Label(ventana_config,          
                        bg= "#C0D4DF",width=180,height=3
                        )
    tit_crear.pack(pady=(0, 20))
    tit_crear.place(x=0,y=0)
    
    #titulo
    imagen_tit = Image.open(abrir_imagen("titulo_secundario.png"))
    imgen_titu = imagen_tit.resize((430, 30))  # Ajusta el tamaño según lo necesario
    imagen_titulo= ImageTk.PhotoImage(imgen_titu)
    tit_crear = tk.Label(ventana_config,
                         
                        font=('Itim', 24, 'bold'),
                        bg= "#C0D4DF",width=440,
                        image=imagen_titulo
                        )
    tit_crear.pack(pady=(0, 20))
    tit_crear.place(x=60,y=8)
    #Titulo 

    # Crear un estilo personalizado para el LabelFrame
    style = ttk.Style()
    style.configure("Custom.TLabelframe", 
                    background=color_secundario,
                    bordercolor="#9DBCCC",
                    relief="groove")
    style.configure("Custom.TLabelframe.Label", 
                    background=color_secundario,
                    foreground="#CACACA")  
    
    # FONDO PARA EL FONDO DE LAS SECCIONES
    imagen_fondo3 = Image.open(abrir_imagen("fondo_secciones.png"))
    imagen_fondo3 = imagen_fondo3.resize((ventana_width, ventana_height))  # Ajustar tamaño si es necesario
    fondo3 = ImageTk.PhotoImage(imagen_fondo3)
    
    #**********************SECCION SUPERIOR********************************  
    seccion_superior= ttk.LabelFrame(ventana_config, text="       ", style="Custom.TLabelframe")
    seccion_superior.pack(padx=0, pady=0, ipadx=0, ipady=0)
    seccion_superior.place(x=60, y=60)


    canvas_central = tk.Canvas(seccion_superior, width=725, height=80, highlightthickness=0,background="#9DBCCC")
    canvas_central.pack(fill="both", expand=True)
    canvas_central.create_image(0, 0, anchor=tk.NW, image=fondo3)
    #para leer la seleccion 
    ttk.Label(ventana_config, text="Selección de archivo PDF", font=('Hero', 12, 'bold'),background="#9DBCCC").pack(anchor=tk.W,padx=(80),pady=(87))
    ruta_pdf = tk.StringVar()
    txt_examinar = tk.Entry(ventana_config, textvariable=ruta_pdf, state='readonly',width=70,font=('Hero',10))#
    txt_examinar.pack(side=tk.LEFT, padx=(0, 10))
    txt_examinar.place(x=80,y=110)
    #cargar imagen para el boton
    imagen_exa = Image.open(abrir_imagen("btnExaminar.png"))
    imagen_ex = imagen_exa.resize((225, 88))  # Ajusta el tamaño según lo necesario
    imagen_e= ImageTk.PhotoImage(imagen_ex)
    btn_examinar = tk.Button(
            ventana_config, 
            padx=20, pady=50,
            width=180,height=56,
            bg="#9DBCCC",
            activebackground="#9DBCCC",
            border=0,
            command=seleccionar_pdf,
            image=imagen_e
            )
    btn_examinar.place(x=590, y=88)
    #********************SECCION CENTRAL*********************************
    #Variables para darle color alos radioButton
    letra_activa = "white"
    color_seleccionado ="white" 
    #
    seccion_central= ttk.LabelFrame(ventana_config, text="       ", style="Custom.TLabelframe")
    seccion_central.pack(padx=0, pady=0, ipadx=0, ipady=0)
    seccion_central.place(x=60, y=155)

    canvas_central = tk.Canvas(seccion_central, width=420, height=294, highlightthickness=0)
    canvas_central.pack(fill="both", expand=True)
    canvas_central.create_image(0, 0, anchor=tk.NW, image=fondo3)
    

    # checkbutton's para TIPO DE BORDE
    seleccion = ttk.Label(canvas_central, text="Selección de tipo de borde ", font=('Hero', 12, 'bold'),background="#9DBCCC")
    canvas_central.create_window(120, 10, window=seleccion)
    
    pie_encabezado = tk.StringVar(value="Encabezado")
    opcion_e = tk.Radiobutton(canvas_central, text="Encabezado", 
                               variable=pie_encabezado, value="Encabezado", 
                               fg='black',  # Color de texto normal
                                bg='#9DBCCC',  # Color de fondo normal
                                activeforeground = letra_activa,  # Color de texto cuando está activo
                                selectcolor= color_seleccionado,  # Color de fondo cuando está seleccionado
                                font=("Times", 12),
                               command=lambda:[opcion_h.config(state=tk.DISABLED),
                                               opcion_v.config(state=tk.DISABLED),
                                               opcion_d.config(state=tk.NORMAL),
                                               opcion_ori.set(0)])
    canvas_central.create_window(86, 40, window=opcion_e)
    
    opcion_p = tk.Radiobutton(canvas_central, text="Pie de página",
                               variable=pie_encabezado, value="Pie de página", 
                                fg='black',  # Color de texto normal
                                bg='#9DBCCC',  # Color de fondo normal
                                activeforeground='white',  # Color de texto cuando está activo
                                selectcolor="white",  # Color de fondo cuando está seleccionado
                                font=("Times", 12),
                               command=lambda:[opcion_h.config(state=tk.NORMAL),
                                               opcion_v.config(state=tk.NORMAL),
                                               opcion_d.config(state=tk.DISABLED)]# Color del texto cuando está habilitado
                        )
    canvas_central.create_window(338, 40, window=opcion_p)
    
    #********************** 
    ##ORIENTACION DE FOLIACION    
    titutlo_secundario = ttk.Label(ventana_config, text="Orientacion de la foliación", font=('Helvetica', 12, 'bold'),background="#9DBCCC")
    opcion_ori = tk.StringVar(value="Horizontal")
    titutlo_secundario.place(x=80,y=240)
    
    opcion_h = tk.Radiobutton(ventana_config, text="Vertical", variable=opcion_ori, value="Vertical",
                          fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12))
    opcion_h.pack(side=tk.LEFT, padx=(0, 20))
    opcion_h.place(x=100, y=270)

    # Opción Horizontal
    opcion_v = tk.Radiobutton(ventana_config, text="Horizontal", variable=opcion_ori, value="Horizontal",
                             fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12))
    opcion_v.pack(side=tk.LEFT)
    opcion_v.place(x=350, y=270)
    
    opcion_h.config(state=tk.DISABLED)
    opcion_v.config(state=tk.DISABLED)
#
    

    # checkbutton's para TIPO DE FOLIACION**************************
    ttk.Label(ventana_config, text="Selección tipo de foliacion", font=('Helvetica', 12, 'bold'),background="#9DBCCC").pack(anchor=tk.W,padx=(77),pady=(114,0))
    opcion_var = tk.StringVar(value="Impar")
    opcion_i = tk.Radiobutton(ventana_config, text="Impar", variable=opcion_var, value="Impar",
                   fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12))
    opcion_i.pack(side=tk.LEFT, padx=(0, 20))
    opcion_i.place(x=100,y=340)
    opcion_n = tk.Radiobutton(ventana_config, text="Normal", 
                   variable=opcion_var, value="Normal",
                   fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12))
    opcion_n.pack(side=tk.LEFT)
    opcion_n.place(x=350,y=340)
    
    imagen_ay = Image.open(abrir_imagen("btnAyuda.png"))
    imagen_ayu = imagen_ay.resize((68, 28))  # Ajusta el tamaño según lo necesario
    imagen_ayuda= ImageTk.PhotoImage(imagen_ayu)
    boton_ayuda = tk.Button(ventana_config, command=mostrar_ayuda,
                            image=imagen_ayuda,width=34,height=28,
                            activebackground="#9DBCCC",bg="#9DBCCC",
                            border=0)
    boton_ayuda.pack(side=tk.RIGHT, padx=5)
    boton_ayuda.place(x=300,y=305)
    #*****************************
    #cargar logo de SUNAT
    imagen_logo = Image.open(abrir_imagen("Logo_Sunat.png"))
    img = imagen_logo.resize((170, 170))  # Ajusta el tamaño según lo necesario
    img_logo= ImageTk.PhotoImage(img)
    logo = tk.Label(ventana_config,image=img_logo,bg=color_secundario)
    logo.pack(pady=(0,30))
    logo.place(x=580,y=245)
    #*****************************
    
#CAMBIAR EL BOTON CREAR
    #  checkbutton's para tipo de ORDEN
    ttk.Label(ventana_config, text="Selección orden", font=('Helvetica', 12, 'bold'),background="#9DBCCC").pack(anchor=tk.W,padx=(77),pady=(40,0))
    opcion_asc_desc = tk.StringVar(value="Ascendente")
    opcion_a = tk.Radiobutton( ventana_config, text="Ascendente", 
                   variable=opcion_asc_desc, value="Ascendente",
                   fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12))
    opcion_a.pack(side=tk.LEFT, padx=(0, 20))
    opcion_a.place(x=100,y=420)
    opcion_d = tk.Radiobutton(ventana_config, text="Descendente",
                   variable=opcion_asc_desc, value="Descendente",
                   fg='black',  # Color de texto normal
                          bg='#9DBCCC',  # Color de fondo normal
                          activeforeground='white',  # Color de texto cuando está activo
                          selectcolor="white",  # Color de fondo cuando está seleccionado
                          font=("Times", 12)
                   )  
    opcion_d.pack(side=tk.LEFT, padx=10, pady=10)  # Añade padding
    opcion_d.place(x=350, y=420)

    #RUTAS PARA GUARDAR (no se muestran, no son necesarios de mostrar)
    ruta_guardar = tk.StringVar()
    aux_ruta_guardar = tk.StringVar()  
    txt_guardar = tk.Entry(ventana_config, textvariable=ruta_guardar, state='readonly', width=50,font=('Hero',12))
    txt_guardar.pack(side=tk.LEFT, padx=(0, 10))
    txt_guardar.place(x=2000,y=2000)

    
    
    #********************************************

    # Crear un marco (frame) para la sección en la ventana secundaria con el estilo personalizado
    seccion_inferior = ttk.LabelFrame(ventana_config, text="       ", style="Custom.TLabelframe")
    seccion_inferior.pack(padx=0, pady=0, ipadx=0, ipady=0)
    seccion_inferior.place(x=60, y=470)

    # Crear un Canvas dentro del LabelFrame para la imagen de fondo
    canvas = tk.Canvas(seccion_inferior, width=725, height=102, highlightthickness=0)
    canvas.pack(fill="both", expand=True)
    canvas.create_image(0, 0, anchor=tk.NW, image=fondo3)

    # Cargar las imágenes de los botones
    imagen_cread = Image.open(abrir_imagen("btnCrearDoc.png"))
    imagen_c = imagen_cread.resize((220, 88))
    imagen_cd = ImageTk.PhotoImage(imagen_c)

    imagen_ca = Image.open(abrir_imagen("btnCancelar.png"))
    imagen_can = imagen_ca.resize((160, 48))
    imagen_cancelar = ImageTk.PhotoImage(imagen_can)

    # Crear los botones dentro del Canvas en el LabelFrame
    boton_crear = tk.Button(canvas, text="", image=imagen_cd,
                            command=crear_doc,width=176, height=55,
                            bg="#9DBCCC", activebackground="#9DBCCC",
                            borderwidth=0, highlightthickness=0)
    canvas.create_window(120, 50, window=boton_crear)

    boton_salir = tk.Button(canvas, image=imagen_cancelar,
                            command=minimizar_ventana_principal,width=177, height=52,
                            bg="#9DBCCC", activebackground="#9DBCCC",
                            borderwidth=0, highlightthickness=0)
    canvas.create_window(618, 51, window=boton_salir)

    #********************************************
    estado = tk.StringVar()
    estado.set("Estado")
    label_estado = tk.Label(ventana_config,textvariable=estado,text="Estado: ",
                            background=color_secundario,font=('Hero', 10, 'bold'))
    label_estado.place(x=80,y=605)
    ventana_config.mainloop()    
#Ventana para convertir docx a PDF
def crear_interfaz():
    from docx2pdf import convert
    #cerrar la ventana principal
    root.iconify()
    #funcion para seleccionar el documento
    def seleccionar_archivo():
        archivo = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx"), ("Word Document", "*.DOCX")])
        if archivo:
            #cambiando de extension, ya que la funcion "convert" no acepta documentos con extensiones mayusculas
            new_file = archivo.replace(".DOCX",".docx")
            try:
                #reemplazamos el archivo con uno nuevo (copia) para que finalmente se guarde en minuscula al extension
                os.replace(archivo, new_file)
                ruta_archivo.set(new_file)
                ventana_convertir.update_idletasks()
            except OSError as e:
                messagebox.showerror(f"Error al renombrar el archivo: {e}")
                      
    #funcion para convertir los  docx
    def convertir_archivo():
        #pedimos la ruta del archivo en una variable
        input = ruta_archivo.get()
        if not input:
            messagebox.showerror("Error", "Seleccione un archivo .docx")
        else:    
            #crea un pdf
            archivo_pdf = os.path.splitext(input)[0] + ".pdf"
            estado.set("Convirtiendo...")
            ventana_convertir.update_idletasks()
            #convierte el archivo a PDF
            convert(input, archivo_pdf,True)
            #mostramos el nombre con el que se guardó el archivo
            estado.set(f"Finalizado: {os.path.basename(archivo_pdf)}")
            messagebox.showinfo("Éxito", f"El archivo ha sido convertido y guardado como: \n{os.path.basename(archivo_pdf)}")
    #minimizar la ventana principal            
    def minimizar():
        root.deiconify()
        ventana_convertir.destroy()
    # Crear la ventana principal
    ventana_convertir = tk.Toplevel(root)
    ventana_convertir.title("Conversor")
    ventana_convertir.geometry("500x250")
    ventana_convertir.minsize(500,250)
    ventana_convertir.maxsize(500,250)
    ventana_convertir.configure(bg=color_secundario)  # Fondo azul claro  
    ventana_convertir.protocol("WM_DELETE_WINDOW",minimizar)
    ruta_archivo = tk.StringVar()
    estado = tk.StringVar()
    # Estilo para los widgets
    estilo = ttk.Style()
    estilo.theme_use('clam')
    estilo.configure('TLabel', background=color_secundario, font=('Arial', 10))
    estilo.configure('TEntry', font=('Arial', 10))
    estilo.configure('TButton', font=('Arial', 10))

    # Título
    titulo = ttk.Label(ventana_convertir, text="Convertir de DOCX a PDF", font=('Arial', 16, 'bold'), background=color_secundario)
    titulo.pack(pady=10)

    # Label e Entry
    label_info = ttk.Label(ventana_convertir, text="Ruta del archivo:")
    label_info.pack(pady=5)

    entry_info = ttk.Entry(ventana_convertir, width=60, state='readonly', textvariable=ruta_archivo)
    entry_info.pack()
    entry_info.insert(0, "Datos de aduana aquí")

    # Frame para los botones
    frame_examinar = ttk.Frame(ventana_convertir)
    frame_examinar.pack(pady=10)

    # Botones
    boton_procesar = ttk.Button(frame_examinar, text="Examinar", style='azul.TButton',command=seleccionar_archivo)
    boton_procesar.pack(side=tk.LEFT, padx=0)

    frame_convertir = ttk.Frame(ventana_convertir)
    frame_convertir.pack(pady=15)
    boton_convertir = ttk.Button(frame_convertir, text="Convertir", style='Rojo.TButton',command=convertir_archivo)
    boton_convertir.pack(side=tk.LEFT, padx=0)

    # Configurar estilos de los botones
    estilo.configure('azul.TButton', background='#096CB7', foreground='white')
    estilo.configure('Rojo.TButton', background='#BD0C4E', foreground='white')
    estilo.map('azul.TButton', background=[('active', '#005596')])
    estilo.map('Rojo.TButton', background=[('active', '#8B0034')])

    # Label de estado
    tk.Label(ventana_convertir, textvariable=estado,bg=color_secundario).pack(pady=10)

    ventana_convertir.mainloop()
# Crear la ventana principal
# Configuración inicial de la ventana
if __name__ == '__main__':
    root = tk.Tk()
    root.title("Herramienta para Foliar Documentos")
    root.geometry("850x640")  # Ajusta el tamaño según sea necesario
    root.maxsize(850,640)
    root.minsize(850,640)
    #paleta de colores

    color_primario = "#D0DDE3"
    color_secundario = "#73A5BC"
    # Cargar la imagen de fondo
    imagen_fondo = Image.open(abrir_imagen("Sunat.jpg"))
    imagen_fondo = imagen_fondo.resize((880,640))  # Ajustar tamaño si es necesario
    fondo = ImageTk.PhotoImage(imagen_fondo)

    # Crear un label para mostrar la imagen de fondo
    label_fondo = tk.Label(root, image=fondo)
    label_fondo.place(x=0, y=0, relwidth=1, relheight=1)
    #LOGO
    imagen_l = Image.open(abrir_imagen("logoSunat.png"))
    imagen_lr = imagen_l.resize((120, 35))  # Ajusta el tamaño según lo necesario
    imagen_logo = ImageTk.PhotoImage(imagen_lr)
    logo = tk.Label(root,image=imagen_logo,bg="#D0DDE3")
    logo.pack(pady=(0,30))
    logo.place(x=170,y=180)

    #TITULOS
    title_label = tk.Label(root, 
                            text="HERRAMIENTA PARA \nFOLIAR DOCUMENTOS",
                            font=('IBM Plex Sans', 28, 'bold'),
                            foreground="black", bg= color_primario
                            )
    title_label.pack(pady=(0, 30))
    title_label.place(x=40,y=260)
    #OGCA    
    title_label = tk.Label(root, 
                        text="O f i c i n a   d e   G e s t i ó n   y   C o o r d i n a c i ó n   A d u a n e r a",
                            font=('Hero', 10,'italic'),
                            foreground="black", bg= color_primario                        
                            )
    title_label.pack(pady=(0, 30))
    title_label.place(x=40,y=380)
    # BOTON CREAR
    imagen_original = Image.open(abrir_imagen("btnCrear.png"))
    imagen_redimensionada = imagen_original.resize((340, 140))  # Ajusta el tamaño según lo necesario
    imagen = ImageTk.PhotoImage(imagen_redimensionada)
    btn_crear_foliar = tk.Button(#   BOTON CREAR
        root, 
        font=("Arial", 14), 
        image=imagen,
        border=0,
        bg=color_primario,#color de fondo
        activebackground=color_primario,
        padx=20, pady=50,
        command=crear_documento_foliado,
        width=305,height=64)
    btn_crear_foliar.place(x=510, y=160)  # Ajusta la posición según sea necesario
    #BOTON CONVERTIR
    #Cargar y redimensionar la imagen
    imagen_convertir = Image.open(abrir_imagen("btnConvertir.png"))
    imagen_c = imagen_convertir.resize((340, 140))  # Ajusta el tamaño según lo necesario
    imagen_cf = ImageTk.PhotoImage(imagen_c)
    btn_convertir_docx = tk.Button(
        root, 
        padx=20, pady=50,
        width=305,height=64,
        bg=color_primario,
        activebackground=color_primario,
        border=0,
        command=crear_interfaz,
        image=imagen_cf)
    btn_convertir_docx.place(x=510, y=300)  # Ajusta la posición según sea necesario
    #BOTON SALIR
    imagen_salir = Image.open(abrir_imagen("btnSalir.png"))
    imagen_r = imagen_salir.resize((340, 140))  # Ajusta el tamaño según lo necesario
    imagen_s = ImageTk.PhotoImage(imagen_r)
    btn_salir = tk.Button(
        root,
        width=292,height=68, 
        bg=color_primario,
        activebackground=color_primario,
        border=0,
        padx=10, pady=50,
        image=imagen_s,
        command=root.destroy)
    btn_salir.place(x=510, y=440)  # Ajusta la posición según sea necesario
    #*************************
    #funcion para ocultar la consola cuando se haga el ejecutable
    hide_console()
    root.mainloop()